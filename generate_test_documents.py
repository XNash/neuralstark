#!/usr/bin/env python3
"""
Generate synthetic test documents for RAG testing
Creates documents with known content for accuracy validation
"""
import os
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from docx import Document
from PIL import Image, ImageDraw, ImageFont
import pandas as pd

# Test content with specific facts for validation
TEST_DOCUMENTS = {
    "financial_report_2024.pdf": {
        "type": "pdf",
        "content": [
            "Financial Report 2024 - NeuralStark Corporation",
            "",
            "Annual Revenue: $5,200,000",
            "Net Profit: $1,150,000",
            "Operating Expenses: $2,800,000",
            "Total Assets: $8,500,000",
            "",
            "Key Performance Indicators:",
            "- Customer Growth Rate: 45% year-over-year",
            "- Market Share: 18% in AI automation sector",
            "- Employee Count: 127 full-time employees",
            "- Product Lines: 5 active product offerings",
            "",
            "Investment Breakdown:",
            "Research & Development: $1,200,000 (23% of revenue)",
            "Sales & Marketing: $950,000 (18% of revenue)",
            "Infrastructure: $650,000 (13% of revenue)"
        ]
    },
    
    "product_catalog_2024.pdf": {
        "type": "pdf",
        "content": [
            "NeuralStark Product Catalog 2024",
            "",
            "Product: AI Document Processor Pro",
            "SKU: AIDP-PRO-2024",
            "Price: $499 per month",
            "Features: OCR, Multi-language support, Cloud storage",
            "",
            "Product: Smart Data Analytics Suite",
            "SKU: SDAS-ENT-2024",
            "Price: $899 per month",
            "Features: Real-time analytics, Custom dashboards, API access",
            "",
            "Product: Automated Workflow Engine",
            "SKU: AWE-STD-2024",
            "Price: $299 per month",
            "Features: Task automation, Integration APIs, Email notifications"
        ]
    },
    
    "employee_handbook.docx": {
        "type": "docx",
        "content": [
            "Employee Handbook 2024",
            "",
            "Company Policies:",
            "",
            "Work Hours: Monday to Friday, 9:00 AM to 5:00 PM",
            "Remote Work: Available 3 days per week",
            "Vacation Days: 25 days per year plus national holidays",
            "Health Insurance: Comprehensive coverage starting day one",
            "",
            "Professional Development:",
            "Annual training budget: $2,000 per employee",
            "Conference attendance: Up to 2 conferences per year",
            "Certification support: Full reimbursement for job-related certifications",
            "",
            "Contact Information:",
            "HR Department: hr@neuralstark.com",
            "IT Support: support@neuralstark.com",
            "Emergency Hotline: +1-555-0199"
        ]
    },
    
    "meeting_notes.docx": {
        "type": "docx",
        "content": [
            "Board Meeting Notes - January 15, 2024",
            "",
            "Attendees: CEO Sarah Johnson, CTO Michael Chen, CFO David Williams",
            "",
            "Key Decisions:",
            "1. Approved budget increase of $500,000 for Q1 2024",
            "2. Launched new client onboarding program starting February 1st",
            "3. Scheduled product launch for AI Chatbot v3.0 on March 15, 2024",
            "",
            "Action Items:",
            "- Michael Chen: Finalize technical specifications by Jan 25",
            "- David Williams: Prepare financial projections by Jan 30",
            "- Sarah Johnson: Schedule client demo sessions for February",
            "",
            "Next Meeting: February 12, 2024 at 10:00 AM"
        ]
    },
    
    "sales_data_q4.xlsx": {
        "type": "xlsx",
        "data": {
            "Month": ["October", "November", "December"],
            "Revenue": [420000, 485000, 520000],
            "New_Customers": [35, 42, 38],
            "Churn_Rate": [3.2, 2.8, 2.5],
            "Average_Deal_Size": [12000, 11500, 13700]
        }
    },
    
    "inventory_status.xlsx": {
        "type": "xlsx",
        "data": {
            "Product_Name": ["Server Rack A", "Laptop Model X", "Monitor 27inch", "Network Switch", "USB Storage"],
            "Quantity": [15, 45, 67, 23, 120],
            "Unit_Price": [2500, 1200, 350, 890, 45],
            "Supplier": ["TechCorp", "CompuMax", "DisplayPro", "NetGear", "StoragePlus"]
        }
    },
    
    "company_overview.txt": {
        "type": "txt",
        "content": """NeuralStark Company Overview

Founded: 2018
Headquarters: San Francisco, California
Industry: Artificial Intelligence & Automation
Mission: Empowering businesses with intelligent automation solutions

Company Size: 127 employees across 5 offices
Global Presence: USA, UK, Germany, Singapore, Australia

Core Technologies:
- Natural Language Processing (NLP)
- Computer Vision and OCR
- Machine Learning Pipelines
- Robotic Process Automation

Notable Achievements:
- Named "Top 50 AI Startups" by TechCrunch 2023
- ISO 27001 Certified for Information Security
- Processed over 10 million documents in 2023
- Client retention rate of 94%

Leadership Team:
CEO: Sarah Johnson (Former VP at Microsoft)
CTO: Michael Chen (PhD in Machine Learning, Stanford)
CFO: David Williams (20 years experience in finance)
"""
    },
    
    "invoice_001.txt": {
        "type": "txt",
        "content": """INVOICE #NS-2024-001

Date: January 20, 2024
Due Date: February 20, 2024

Bill To:
Acme Corporation
123 Business Street
New York, NY 10001

Services Provided:
- AI Document Processing (Dec 2023): $499.00
- Cloud Storage (50GB): $79.00
- API Access Premium: $199.00
- Technical Support: $150.00

Subtotal: $927.00
Tax (8%): $74.16
Total Due: $1,001.16

Payment Methods:
Bank Transfer: Account #1234567890
Credit Card: Visa/Mastercard/Amex accepted
Wire Transfer: Swift Code NEUSTARK01

Thank you for your business!
"""
    }
}

def create_pdf_document(filename, content_lines):
    """Create a PDF document with given content"""
    doc = SimpleDocTemplate(filename, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    for line in content_lines:
        if line.strip():
            if line.endswith(":") or line.startswith("Product:") or line.startswith("SKU:"):
                story.append(Paragraph(f"<b>{line}</b>", styles['Normal']))
            else:
                story.append(Paragraph(line, styles['Normal']))
            story.append(Spacer(1, 0.1 * 72))
        else:
            story.append(Spacer(1, 0.2 * 72))
    
    doc.build(story)
    print(f"✓ Created: {filename}")

def create_docx_document(filename, content_lines):
    """Create a DOCX document with given content"""
    doc = Document()
    
    for line in content_lines:
        if line.strip():
            if line.endswith(":") and not line.startswith(" "):
                # Heading
                doc.add_heading(line, level=2)
            else:
                doc.add_paragraph(line)
        else:
            doc.add_paragraph()
    
    doc.save(filename)
    print(f"✓ Created: {filename}")

def create_xlsx_document(filename, data_dict):
    """Create an XLSX spreadsheet with given data"""
    df = pd.DataFrame(data_dict)
    df.to_excel(filename, index=False, engine='openpyxl')
    print(f"✓ Created: {filename}")

def create_txt_document(filename, content):
    """Create a text document"""
    with open(filename, 'w', encoding='utf-8') as f:
        f.write(content)
    print(f"✓ Created: {filename}")

def create_ocr_image(filename, text_lines):
    """Create an image with text for OCR testing"""
    # Create a white image
    img = Image.new('RGB', (800, 600), color='white')
    draw = ImageDraw.Draw(img)
    
    # Try to use a default font, fallback to default if not available
    try:
        font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 20)
    except:
        font = ImageFont.load_default()
    
    # Draw text
    y_position = 50
    for line in text_lines:
        draw.text((50, y_position), line, fill='black', font=font)
        y_position += 35
    
    img.save(filename)
    print(f"✓ Created: {filename}")

def generate_all_documents():
    """Generate all test documents"""
    print("=" * 60)
    print("Generating Synthetic Test Documents for RAG Testing")
    print("=" * 60)
    print()
    
    # Create knowledge base directory
    kb_path = "/app/backend/knowledge_base/internal"
    os.makedirs(kb_path, exist_ok=True)
    
    # Generate each document
    for filename, doc_info in TEST_DOCUMENTS.items():
        filepath = os.path.join(kb_path, filename)
        
        if doc_info["type"] == "pdf":
            create_pdf_document(filepath, doc_info["content"])
        elif doc_info["type"] == "docx":
            create_docx_document(filepath, doc_info["content"])
        elif doc_info["type"] == "xlsx":
            create_xlsx_document(filepath, doc_info["data"])
        elif doc_info["type"] == "txt":
            create_txt_document(filepath, doc_info["content"])
    
    # Create OCR test images
    print()
    print("Creating OCR test images...")
    
    ocr_image_1 = os.path.join(kb_path, "contact_info.png")
    create_ocr_image(ocr_image_1, [
        "NeuralStark Contact Information",
        "",
        "Main Office: +1-555-0123",
        "Sales: sales@neuralstark.com",
        "Support: support@neuralstark.com",
        "",
        "Address:",
        "456 Innovation Drive",
        "San Francisco, CA 94105"
    ])
    
    ocr_image_2 = os.path.join(kb_path, "pricing_table.png")
    create_ocr_image(ocr_image_2, [
        "Pricing Plans 2024",
        "",
        "Starter Plan: $199/month",
        "Professional Plan: $499/month",
        "Enterprise Plan: $999/month",
        "",
        "All plans include:",
        "- 24/7 Support",
        "- Free Updates",
        "- 99.9% Uptime SLA"
    ])
    
    print()
    print("=" * 60)
    print(f"✓ Successfully generated {len(TEST_DOCUMENTS) + 2} test documents")
    print(f"Location: {kb_path}")
    print("=" * 60)
    print()
    print("Documents created:")
    print("- 2 PDF files (financial report, product catalog)")
    print("- 2 DOCX files (employee handbook, meeting notes)")
    print("- 2 XLSX files (sales data, inventory)")
    print("- 2 TXT files (company overview, invoice)")
    print("- 2 PNG images for OCR testing")
    print()

if __name__ == "__main__":
    generate_all_documents()
