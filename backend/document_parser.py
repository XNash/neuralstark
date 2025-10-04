import os
from typing import Optional
import logging
import pandas as pd
import json
import subprocess
import tempfile
import shutil
from pathlib import Path

# Document parsing libraries
try:
    import pypdf
except ImportError:
    pypdf = None
    logging.warning("pypdf not installed. PDF parsing will not be available.")

try:
    import docx
except ImportError:
    docx = None
    logging.warning("python-docx not installed. DOCX parsing will not be available.")

try:
    from odfdo import Document
except ImportError:
    Document = None
    logging.warning("odfdo not installed. ODT parsing will not be available.")

# OCR libraries
try:
    from PIL import Image
    import pytesseract
    from pdf2image import convert_from_path
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    logging.warning("OCR libraries not installed. OCR functionality will not be available.")

logging.basicConfig(level=logging.INFO,
                    format='%(asctime)s - %(message)s',
                    datefmt='%Y-%m-%d %H:%M:%S')

def perform_ocr(image_path: str, languages: str = "eng+fra") -> Optional[str]:
    """Performs OCR on an image file."""
    if not OCR_AVAILABLE:
        logging.warning("OCR not available. Skipping OCR.")
        return None
    
    try:
        image = Image.open(image_path)
        text = pytesseract.image_to_string(image, lang=languages)
        return text.strip()
    except Exception as e:
        logging.error(f"Error performing OCR on {image_path}: {e}")
        return None

def perform_ocr_on_image_object(image, languages: str = "eng+fra") -> Optional[str]:
    """Performs OCR on a PIL Image object."""
    if not OCR_AVAILABLE:
        return None
    
    try:
        text = pytesseract.image_to_string(image, lang=languages)
        return text.strip()
    except Exception as e:
        logging.error(f"Error performing OCR on image object: {e}")
        return None

def convert_doc_to_docx(doc_path: str) -> Optional[str]:
    """Converts a .doc file to .docx using LibreOffice."""
    try:
        # Create a temporary directory for conversion
        temp_dir = tempfile.mkdtemp()
        
        # Use LibreOffice to convert
        cmd = [
            'libreoffice',
            '--headless',
            '--convert-to', 'docx',
            '--outdir', temp_dir,
            doc_path
        ]
        
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
        
        if result.returncode != 0:
            logging.error(f"LibreOffice conversion failed: {result.stderr}")
            shutil.rmtree(temp_dir)
            return None
        
        # Find the converted file
        doc_name = Path(doc_path).stem
        docx_path = os.path.join(temp_dir, f"{doc_name}.docx")
        
        if os.path.exists(docx_path):
            return docx_path
        else:
            logging.error(f"Converted file not found: {docx_path}")
            shutil.rmtree(temp_dir)
            return None
            
    except Exception as e:
        logging.error(f"Error converting .doc to .docx: {e}")
        return None

def parse_text_file(file_path: str) -> Optional[str]:
    """Parses a plain text file."""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        logging.error(f"Error parsing text file {file_path}: {e}")
        return None

def parse_pdf_file(file_path: str, ocr_enabled: bool = True) -> Optional[str]:
    """Parses a PDF file with OCR support for images and scanned pages."""
    if not pypdf:
        logging.error("pypdf is not installed. Cannot parse PDF files.")
        return None
    
    try:
        text = ""
        
        # First, try standard text extraction
        with open(file_path, 'rb') as f:
            reader = pypdf.PdfReader(f)
            for page_num, page in enumerate(reader.pages):
                page_text = page.extract_text() or ""
                text += page_text
                
                # If page has very little text, it might be scanned - try OCR
                if ocr_enabled and OCR_AVAILABLE and len(page_text.strip()) < 50:
                    logging.info(f"Page {page_num + 1} has minimal text. Attempting OCR...")
                    try:
                        # Convert PDF page to image and OCR
                        images = convert_from_path(file_path, first_page=page_num+1, last_page=page_num+1)
                        if images:
                            ocr_text = perform_ocr_on_image_object(images[0])
                            if ocr_text:
                                text += f"\n[OCR from page {page_num + 1}]\n{ocr_text}\n"
                    except Exception as e:
                        logging.warning(f"OCR failed for page {page_num + 1}: {e}")
        
        # Extract and OCR embedded images
        if ocr_enabled and OCR_AVAILABLE:
            try:
                with open(file_path, 'rb') as f:
                    reader = pypdf.PdfReader(f)
                    for page_num, page in enumerate(reader.pages):
                        if '/XObject' in page['/Resources']:
                            xObject = page['/Resources']['/XObject'].get_object()
                            
                            for obj in xObject:
                                if xObject[obj]['/Subtype'] == '/Image':
                                    try:
                                        # Extract image and perform OCR
                                        size = (xObject[obj]['/Width'], xObject[obj]['/Height'])
                                        data = xObject[obj].get_data()
                                        
                                        # Create PIL Image from data
                                        mode = "RGB" if xObject[obj]['/ColorSpace'] == '/DeviceRGB' else "L"
                                        image = Image.frombytes(mode, size, data)
                                        
                                        ocr_text = perform_ocr_on_image_object(image)
                                        if ocr_text:
                                            text += f"\n[OCR from embedded image in page {page_num + 1}]\n{ocr_text}\n"
                                    except Exception as e:
                                        logging.debug(f"Could not OCR image in page {page_num + 1}: {e}")
            except Exception as e:
                logging.warning(f"Error extracting images from PDF: {e}")
        
        return text if text.strip() else None
    except Exception as e:
        logging.error(f"Error parsing PDF file {file_path}: {e}")
        return None

def parse_docx_file(file_path: str, ocr_enabled: bool = True) -> Optional[str]:
    """Parses a DOCX file with OCR support for embedded images."""
    if not docx:
        logging.error("python-docx is not installed. Cannot parse DOCX files.")
        return None
    
    try:
        document = docx.Document(file_path)
        text = "\n".join([paragraph.text for paragraph in document.paragraphs])
        
        # Extract and OCR images if enabled
        if ocr_enabled and OCR_AVAILABLE:
            try:
                for rel in document.part.rels.values():
                    if "image" in rel.target_ref:
                        try:
                            image_data = rel.target_part.blob
                            image = Image.open(io.BytesIO(image_data))
                            ocr_text = perform_ocr_on_image_object(image)
                            if ocr_text:
                                text += f"\n[OCR from embedded image]\n{ocr_text}\n"
                        except Exception as e:
                            logging.debug(f"Could not OCR embedded image: {e}")
            except Exception as e:
                logging.warning(f"Error extracting images from DOCX: {e}")
        
        return text
    except Exception as e:
        logging.error(f"Error parsing DOCX file {file_path}: {e}")
        return None

def parse_doc_file(file_path: str, ocr_enabled: bool = True) -> Optional[str]:
    """Parses a .doc file by converting it to .docx first using LibreOffice."""
    try:
        logging.info(f"Converting .doc file to .docx: {file_path}")
        docx_path = convert_doc_to_docx(file_path)
        
        if docx_path:
            text = parse_docx_file(docx_path, ocr_enabled)
            
            # Clean up temporary file
            try:
                temp_dir = os.path.dirname(docx_path)
                shutil.rmtree(temp_dir)
            except Exception as e:
                logging.warning(f"Could not clean up temporary directory: {e}")
            
            return text
        else:
            logging.error(f"Failed to convert .doc file: {file_path}")
            return None
    except Exception as e:
        logging.error(f"Error parsing .doc file {file_path}: {e}")
        return None

def parse_image_file(file_path: str) -> Optional[str]:
    """Parses an image file using OCR."""
    if not OCR_AVAILABLE:
        logging.error("OCR libraries not installed. Cannot parse image files.")
        return None
    
    try:
        text = perform_ocr(file_path)
        return text if text else None
    except Exception as e:
        logging.error(f"Error parsing image file {file_path}: {e}")
        return None

def parse_csv_file(file_path: str) -> Optional[str]:
    """Parses a CSV file and returns its content as a string (e.g., Markdown table)."""
    try:
        df = pd.read_csv(file_path)
        return df.to_markdown(index=False)
    except Exception as e:
        logging.error(f"Error parsing CSV file {file_path}: {e}")
        return None

def parse_excel_file(file_path: str) -> Optional[str]:
    """Parses an Excel file (XLSX/XLS) and returns content of all sheets as a string."""
    try:
        xls = pd.ExcelFile(file_path)
        all_sheets_content = []
        for sheet_name in xls.sheet_names:
            df = pd.read_excel(xls, sheet_name=sheet_name)
            all_sheets_content.append(f"\n--- Sheet: {sheet_name} ---\n")
            all_sheets_content.append(df.to_markdown(index=False))
        return "\n".join(all_sheets_content)
    except Exception as e:
        logging.error(f"Error parsing Excel file {file_path}: {e}")
        return None

def parse_markdown_file(file_path: str) -> Optional[str]:
    """Parses a Markdown file as plain text."""
    return parse_text_file(file_path)

def parse_json_file(file_path: str) -> Optional[str]:
    """Parses a JSON file and returns its content as a formatted string."""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
            return json.dumps(data, indent=2)
    except Exception as e:
        logging.error(f"Error parsing JSON file {file_path}: {e}")
        return None

def parse_odt_file(file_path: str) -> Optional[str]:
    """Parses an ODT file and returns its text content."""
    if not Document:
        logging.error("odfdo not installed. Cannot parse ODT files.")
        return None
    try:
        doc = Document(file_path)
        return doc.get_text()
    except Exception as e:
        logging.error(f"Error parsing ODT file {file_path}: {e}")
        return None

def parse_document(file_path: str, ocr_enabled: bool = True) -> Optional[str]:
    """Parses a document based on its file extension with OCR support."""
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()

    # Get OCR setting from environment if available
    from backend.config import settings
    if hasattr(settings, 'OCR_ENABLED'):
        ocr_enabled = ocr_enabled and settings.OCR_ENABLED

    if ext == '.txt':
        return parse_text_file(file_path)
    elif ext == '.pdf':
        return parse_pdf_file(file_path, ocr_enabled)
    elif ext == '.docx':
        return parse_docx_file(file_path, ocr_enabled)
    elif ext == '.doc':
        return parse_doc_file(file_path, ocr_enabled)
    elif ext == '.csv':
        return parse_csv_file(file_path)
    elif ext in ['.xls', '.xlsx']:
        return parse_excel_file(file_path)
    elif ext == '.md':
        return parse_markdown_file(file_path)
    elif ext == '.json':
        return parse_json_file(file_path)
    elif ext == '.odt':
        return parse_odt_file(file_path)
    elif ext in ['.png', '.jpg', '.jpeg', '.tiff', '.tif', '.bmp', '.gif']:
        return parse_image_file(file_path)
    else:
        logging.warning(f"Unsupported file type: {ext} for {file_path}")
        return None
