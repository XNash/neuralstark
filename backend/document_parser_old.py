import os
from typing import Optional
import logging
import pandas as pd
import json # Added for JSON parsing

# Document parsing libraries
try:
    import pypdf
except ImportError:
    pypdf = None
    logging.warning("pypdf not installed. PDF parsing will not be available.")

try:
    import docx
except ImportError:
    docx = None
    logging.warning("python-docx not installed. DOCX parsing will not be available.")

try:
    from odfdo import Document
except ImportError:
    Document = None
    logging.warning("odfdo not installed. ODT parsing will not be available.")

logging.basicConfig(level=logging.INFO,
                    format='%(asctime)s - %(message)s',
                    datefmt='%Y-%m-%d %H:%M:%S')

def parse_text_file(file_path: str) -> Optional[str]:
    """Parses a plain text file."""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        logging.error(f"Error parsing text file {file_path}: {e}")
        return None

def parse_pdf_file(file_path: str) -> Optional[str]:
    """Parses a PDF file."""
    if not pypdf:
        logging.error("pypdf is not installed. Cannot parse PDF files.")
        return None
    try:
        text = ""
        with open(file_path, 'rb') as f:
            reader = pypdf.PdfReader(f)
            for page in reader.pages:
                text += page.extract_text() or ""
        return text
    except Exception as e:
        logging.error(f"Error parsing PDF file {file_path}: {e}")
        return None

def parse_docx_file(file_path: str) -> Optional[str]:
    """Parses a DOCX file."""
    if not docx:
        logging.error("python-docx is not installed. Cannot parse DOCX files.")
        return None
    try:
        document = docx.Document(file_path)
        return "\n".join([paragraph.text for paragraph in document.paragraphs])
    except Exception as e:
        logging.error(f"Error parsing DOCX file {file_path}: {e}")
        return None

def parse_csv_file(file_path: str) -> Optional[str]:
    """Parses a CSV file and returns its content as a string (e.g., Markdown table)."""
    try:
        df = pd.read_csv(file_path)
        # Convert DataFrame to a string representation, e.g., Markdown table
        return df.to_markdown(index=False)
    except Exception as e:
        logging.error(f"Error parsing CSV file {file_path}: {e}")
        return None

def parse_excel_file(file_path: str) -> Optional[str]:
    """Parses an Excel file (XLSX/XLS) and returns content of all sheets as a string."""
    try:
        xls = pd.ExcelFile(file_path)
        all_sheets_content = []
        for sheet_name in xls.sheet_names:
            df = pd.read_excel(xls, sheet_name=sheet_name)
            all_sheets_content.append(f"\n--- Sheet: {sheet_name} ---\n")
            all_sheets_content.append(df.to_markdown(index=False))
        return "\n".join(all_sheets_content)
    except Exception as e:
        logging.error(f"Error parsing Excel file {file_path}: {e}")
        return None

def parse_markdown_file(file_path: str) -> Optional[str]:
    """Parses a Markdown file as plain text."""
    return parse_text_file(file_path) # Markdown can be treated as plain text for RAG

def parse_json_file(file_path: str) -> Optional[str]:
    """Parses a JSON file and returns its content as a formatted string."""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
            return json.dumps(data, indent=2) # Pretty print JSON
    except Exception as e:
        logging.error(f"Error parsing JSON file {file_path}: {e}")
        return None

def parse_odt_file(file_path: str) -> Optional[str]:
    """Parses an ODT file and returns its text content."""
    if not Document:
        logging.error("odfdo not installed. Cannot parse ODT files.")
        return None
    try:
        doc = Document(file_path)
        return doc.get_text()
    except Exception as e:
        logging.error(f"Error parsing ODT file {file_path}: {e}")
        return None

def parse_document(file_path: str) -> Optional[str]:
    """Parses a document based on its file extension."""
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()

    if ext == '.txt':
        return parse_text_file(file_path)
    elif ext == '.pdf':
        return parse_pdf_file(file_path)
    elif ext == '.docx':
        return parse_docx_file(file_path)
    elif ext == '.csv':
        return parse_csv_file(file_path)
    elif ext in ['.xls', '.xlsx']:
        return parse_excel_file(file_path)
    elif ext == '.md':
        return parse_markdown_file(file_path)
    elif ext == '.json':
        return parse_json_file(file_path)
    elif ext == '.odt':
        return parse_odt_file(file_path)
    else:
        logging.warning(f"Unsupported file type: {ext} for {file_path}")
        return None
