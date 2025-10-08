#!/usr/bin/env python3
"""
Create new test documents with specific content for RAG testing
"""

import os
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from PIL import Image, ImageDraw, ImageFont
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DIR = "/app/backend/knowledge_base/internal"

def create_sales_report_xlsx():
    """Create a detailed sales report with Q4 2024 data"""
    filename = os.path.join(OUTPUT_DIR, "sales_report_q4_2024.xlsx")
    
    # Monthly sales data
    monthly_sales = {
        'Month': ['October', 'November', 'December'],
        'North_Region': [125000, 142000, 198000],
        'South_Region': [98000, 105000, 156000],
        'East_Region': [87000, 91000, 134000],
        'West_Region': [110000, 128000, 175000],
        'Total_Sales': [420000, 466000, 663000]
    }
    
    # Product performance
    products = {
        'Product_ID': ['P001', 'P002', 'P003', 'P004', 'P005'],
        'Product_Name': ['SmartWatch Pro', 'Wireless Earbuds', 'Fitness Tracker', 'Smart Scale', 'Sleep Monitor'],
        'Units_Sold': [2450, 5620, 3890, 1230, 890],
        'Revenue': [367500, 280900, 194500, 73800, 53400],
        'Profit_Margin': ['18%', '25%', '22%', '15%', '20%']
    }
    
    # Sales rep performance
    sales_reps = {
        'Rep_Name': ['Michael Chen', 'Sarah Williams', 'David Martinez', 'Emma Thompson', 'James Brown'],
        'Region': ['North', 'South', 'East', 'West', 'North'],
        'Total_Sales': [285000, 245000, 198000, 267000, 180000],
        'Commission': [14250, 12250, 9900, 13350, 9000],
        'Customer_Rating': [4.8, 4.9, 4.6, 4.7, 4.5]
    }
    
    with pd.ExcelWriter(filename, engine='openpyxl') as writer:
        pd.DataFrame(monthly_sales).to_excel(writer, sheet_name='Monthly_Sales', index=False)
        pd.DataFrame(products).to_excel(writer, sheet_name='Product_Performance', index=False)
        pd.DataFrame(sales_reps).to_excel(writer, sheet_name='Sales_Reps', index=False)
        
        # Summary sheet
        summary = {
            'Metric': ['Total Q4 Sales', 'Average Monthly Sales', 'Top Product', 'Top Sales Rep', 'Best Month'],
            'Value': ['$1,549,000', '$516,333', 'Wireless Earbuds', 'Michael Chen', 'December']
        }
        pd.DataFrame(summary).to_excel(writer, sheet_name='Summary', index=False)
    
    print(f"✓ Created: sales_report_q4_2024.xlsx")
    return filename

def create_project_proposal_docx():
    """Create a project proposal document"""
    filename = os.path.join(OUTPUT_DIR, "ai_chatbot_project_proposal.docx")
    
    doc = Document()
    
    # Title
    title = doc.add_heading('AI Customer Service Chatbot - Project Proposal', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Executive Summary
    doc.add_heading('Executive Summary', 1)
    doc.add_paragraph(
        'This proposal outlines the development of an AI-powered customer service chatbot '
        'for GlobalTech Solutions. The chatbot will utilize advanced natural language processing '
        'to handle customer inquiries 24/7, reducing response time by 85% and support costs by 60%.'
    )
    
    # Project Details
    doc.add_heading('Project Scope', 1)
    doc.add_paragraph('Project Name: GlobalTech AI Assistant')
    doc.add_paragraph('Project Manager: Rebecca Foster')
    doc.add_paragraph('Budget: $450,000')
    doc.add_paragraph('Timeline: 6 months (January 2025 - June 2025)')
    doc.add_paragraph('Team Size: 8 members')
    
    # Technical Stack
    doc.add_heading('Technical Architecture', 1)
    tech_items = [
        'Frontend: React.js with TypeScript',
        'Backend: Python FastAPI',
        'AI Model: GPT-4 Turbo with custom fine-tuning',
        'Vector Database: Pinecone for knowledge retrieval',
        'Deployment: AWS with auto-scaling',
        'Security: End-to-end encryption, SOC 2 compliance'
    ]
    for item in tech_items:
        doc.add_paragraph(item, style='List Bullet')
    
    # Deliverables
    doc.add_heading('Key Deliverables', 1)
    deliverables = [
        'Phase 1 (Month 1-2): Requirements gathering and system design',
        'Phase 2 (Month 3-4): AI model development and training',
        'Phase 3 (Month 5): Integration and testing',
        'Phase 4 (Month 6): Deployment and training'
    ]
    for item in deliverables:
        doc.add_paragraph(item, style='List Number')
    
    # Expected Outcomes
    doc.add_heading('Expected Outcomes', 1)
    doc.add_paragraph(
        'The AI chatbot is expected to handle 75% of customer inquiries autonomously, '
        'with an average response time of 3 seconds. Customer satisfaction scores are '
        'projected to increase from 78% to 92% within the first quarter of deployment.'
    )
    
    # Contact
    doc.add_heading('Contact Information', 1)
    doc.add_paragraph('Email: rebecca.foster@globaltech.com')
    doc.add_paragraph('Phone: +1-555-0199')
    doc.add_paragraph('Department: Innovation & AI')
    
    doc.save(filename)
    print(f"✓ Created: ai_chatbot_project_proposal.docx")
    return filename

def create_legal_contract_pdf():
    """Create a legal contract PDF"""
    filename = os.path.join(OUTPUT_DIR, "software_license_agreement.pdf")
    
    doc = SimpleDocTemplate(filename, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    # Title
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        textColor=colors.HexColor('#000000'),
        spaceAfter=30,
        alignment=1  # Center
    )
    story.append(Paragraph("SOFTWARE LICENSE AGREEMENT", title_style))
    story.append(Spacer(1, 0.3*inch))
    
    # Agreement details
    story.append(Paragraph("Agreement Number: SLA-2024-7856", styles['Normal']))
    story.append(Paragraph("Effective Date: November 1, 2024", styles['Normal']))
    story.append(Paragraph("License Type: Enterprise Multi-User", styles['Normal']))
    story.append(Spacer(1, 0.2*inch))
    
    # Parties
    story.append(Paragraph("PARTIES TO THIS AGREEMENT", styles['Heading2']))
    story.append(Paragraph(
        "<b>Licensor:</b> TechVision Software Inc.<br/>"
        "Address: 789 Innovation Drive, Silicon Valley, CA 94025<br/>"
        "Contact: licensing@techvision.com",
        styles['Normal']
    ))
    story.append(Spacer(1, 0.1*inch))
    story.append(Paragraph(
        "<b>Licensee:</b> DataCorp Industries<br/>"
        "Address: 456 Business Plaza, New York, NY 10001<br/>"
        "Contact: legal@datacorp.com",
        styles['Normal']
    ))
    story.append(Spacer(1, 0.2*inch))
    
    # License terms
    story.append(Paragraph("TERMS AND CONDITIONS", styles['Heading2']))
    
    terms = [
        "<b>1. Grant of License:</b> TechVision hereby grants DataCorp a non-exclusive, "
        "non-transferable license to use the Analytics Pro software suite.",
        
        "<b>2. License Fee:</b> Total license fee is $125,000 USD, payable in three installments: "
        "$50,000 upon signing, $40,000 after 3 months, and $35,000 after 6 months.",
        
        "<b>3. User Limit:</b> This license permits up to 500 concurrent users. Additional user "
        "licenses can be purchased at $200 per user per year.",
        
        "<b>4. Support Services:</b> TechVision will provide 24/7 technical support and quarterly "
        "software updates for the duration of this agreement.",
        
        "<b>5. Term:</b> This agreement is valid for 3 years from the effective date and will "
        "auto-renew unless terminated with 90 days written notice.",
        
        "<b>6. Termination Clause:</b> Either party may terminate this agreement with 90 days "
        "written notice. Early termination penalty is 25% of remaining license value.",
    ]
    
    for term in terms:
        story.append(Paragraph(term, styles['Normal']))
        story.append(Spacer(1, 0.15*inch))
    
    # Signatures
    story.append(Spacer(1, 0.3*inch))
    story.append(Paragraph("SIGNATURES", styles['Heading2']))
    
    signature_data = [
        ['Licensor Representative:', 'Licensee Representative:'],
        ['Name: Jennifer Kim', 'Name: Robert Anderson'],
        ['Title: VP of Licensing', 'Title: CTO'],
        ['Date: November 1, 2024', 'Date: November 1, 2024'],
    ]
    
    sig_table = Table(signature_data, colWidths=[3*inch, 3*inch])
    sig_table.setStyle(TableStyle([
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
    ]))
    story.append(sig_table)
    
    doc.build(story)
    print(f"✓ Created: software_license_agreement.pdf")
    return filename

def create_scanned_receipt_pdf_ocr():
    """Create a scanned receipt image that looks like a PDF scan"""
    # First create the receipt image
    img_filename = os.path.join(OUTPUT_DIR, "restaurant_receipt_scan.png")
    
    img = Image.new('RGB', (600, 900), color='white')
    draw = ImageDraw.Draw(img)
    
    try:
        font_large = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 28)
        font_medium = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 22)
        font_small = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 18)
    except:
        font_large = ImageFont.load_default()
        font_medium = ImageFont.load_default()
        font_small = ImageFont.load_default()
    
    # Receipt header
    draw.text((200, 30), "THE BISTRO", fill='black', font=font_large)
    draw.text((150, 70), "123 Main Street, Chicago, IL", fill='black', font=font_small)
    draw.text((180, 95), "Phone: (312) 555-0147", fill='black', font=font_small)
    
    draw.line((50, 120, 550, 120), fill='black', width=2)
    
    # Receipt details
    draw.text((50, 140), "Receipt #: R-2024-08956", fill='black', font=font_medium)
    draw.text((50, 170), "Date: October 15, 2024", fill='black', font=font_medium)
    draw.text((50, 200), "Time: 7:45 PM", fill='black', font=font_medium)
    draw.text((50, 230), "Server: Maria Rodriguez", fill='black', font=font_medium)
    draw.text((50, 260), "Table: 12", fill='black', font=font_medium)
    
    draw.line((50, 290, 550, 290), fill='black', width=2)
    
    # Items
    draw.text((50, 310), "ITEMS:", fill='black', font=font_medium)
    
    items = [
        ("Grilled Salmon", "$28.50"),
        ("Caesar Salad", "$12.00"),
        ("Pasta Carbonara", "$22.00"),
        ("Tiramisu (x2)", "$16.00"),
        ("House Red Wine", "$35.00"),
        ("Sparkling Water", "$6.00"),
    ]
    
    y = 350
    for item, price in items:
        draw.text((70, y), item, fill='black', font=font_small)
        draw.text((450, y), price, fill='black', font=font_small)
        y += 35
    
    draw.line((50, y + 10, 550, y + 10), fill='black', width=1)
    
    # Totals
    y += 30
    draw.text((70, y), "Subtotal:", fill='black', font=font_medium)
    draw.text((450, y), "$119.50", fill='black', font=font_medium)
    y += 35
    draw.text((70, y), "Tax (8.5%):", fill='black', font=font_medium)
    draw.text((450, y), "$10.16", fill='black', font=font_medium)
    y += 35
    draw.text((70, y), "Tip (18%):", fill='black', font=font_medium)
    draw.text((450, y), "$21.51", fill='black', font=font_medium)
    y += 35
    
    draw.line((50, y + 5, 550, y + 5), fill='black', width=2)
    y += 25
    
    draw.text((70, y), "TOTAL:", fill='black', font=font_large)
    draw.text((430, y), "$151.17", fill='black', font=font_large)
    
    y += 60
    draw.text((150, y), "Payment Method: Visa **** 4532", fill='black', font=font_small)
    y += 30
    draw.text((180, y), "Thank you for dining with us!", fill='black', font=font_small)
    
    img.save(img_filename)
    print(f"✓ Created: restaurant_receipt_scan.png (for OCR testing)")
    
    # Now create a PDF version
    pdf_filename = os.path.join(OUTPUT_DIR, "restaurant_receipt_scan.pdf")
    from reportlab.pdfgen import canvas
    
    c = canvas.Canvas(pdf_filename, pagesize=letter)
    c.drawImage(img_filename, 0, 0, width=6*inch, height=9*inch)
    c.save()
    
    print(f"✓ Created: restaurant_receipt_scan.pdf (scanned image as PDF)")
    return pdf_filename

def main():
    print("="*70)
    print("  Creating New Test Documents for RAG Testing")
    print("="*70)
    print()
    
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    
    # Create documents
    files_created = []
    
    files_created.append(create_sales_report_xlsx())
    files_created.append(create_project_proposal_docx())
    files_created.append(create_legal_contract_pdf())
    files_created.append(create_scanned_receipt_pdf_ocr())
    
    print()
    print("="*70)
    print("  ✅ All new test documents created!")
    print("="*70)
    print("\nFiles created:")
    for f in files_created:
        print(f"  • {os.path.basename(f)}")
    print()
    
    return files_created

if __name__ == "__main__":
    main()
